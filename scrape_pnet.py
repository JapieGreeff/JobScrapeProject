""" scraping from PNet advanced search urls to create job listings """
import urllib.request as scraper
from bs4 import BeautifulSoup
from job_listing import JobListing
from docx import Document
import time
import os
import pickle
import logging
from datetime import datetime

#This is the old styling on PNet - they seem to have changed it, maybe in response to my scrape.
def scrapePrelimListings(counter, listingPage):
    """ takes in a advanced search listing page and extracts all of the job listings from it. counter will be used to generate ids for listings"""
    # the page you input will be the url of a search that is full of listings. set the results page to show 100 results at a time for maximum impact
    #listingPage = "https://www.pnet.co.za/5/job-search-detailed.html?searchTypeFrom=detailedSearch&searchOrigin=Detailed-Search_detailed-search&newsearch=1&Function=1003001&Function=1003006&Function=1010000&Country=1000204&Country=1000205&Country=1000206&Country=1000207&Country=1000208&Country=1000209&keyword=Python"
    page = scraper.urlopen(listingPage) 
    soup = BeautifulSoup(page, features="lxml")
    positionsPreTextScrape = []
    #articleBodies = soup.find_all('article', class_='job-element')
    articleBodies = soup.find_all('article')
    tracking_number = counter
    for article in articleBodies:
        body = article.find('div', class_='job-element__body')
        url = body.find('a', class_='job-element__url', href=True)
        company = body.find('div', class_='job-element__body__company')
        location = body.find('li', class_='job-element__body__location')
        details = body.find('div', class_='job-element__body__details')
        detailTable = details.find('dl', class_='job-element__body__meta')
        detailsTitles = detailTable.find_all('dt', class_='job-element__body__meta-title')
        detailsValues = detailTable.find_all('dd', class_='job-element__body__meta-value')
        if url['href']:
            #print(f'{url["href"]} {url.text.strip()} {company.text.strip()} ')
            detailsText = []
            for detail in zip(detailsTitles, detailsValues):
                detailsText.append((detail[0].text, detail[1].text))
            newListing = JobListing(f"JL{str(tracking_number)}", url["href"], company.text.strip(), url.text.strip(), detailsText, location.text.strip())
            positionsPreTextScrape.append(newListing)
            tracking_number = tracking_number + 1
            logging.info(newListing)
    return positionsPreTextScrape, tracking_number

def scrapePrelimListingsNotClasses(counter, listingPage):
    """ takes in a advanced search listing page and extracts all of the job listings from it. counter will be used to generate ids for listings"""
    # the page you input will be the url of a search that is full of listings. set the results page to show 100 results at a time for maximum impact
    #listingPage = "https://www.pnet.co.za/5/job-search-detailed.html?searchTypeFrom=detailedSearch&searchOrigin=Detailed-Search_detailed-search&newsearch=1&Function=1003001&Function=1003006&Function=1010000&Country=1000204&Country=1000205&Country=1000206&Country=1000207&Country=1000208&Country=1000209&keyword=Python"
    page = scraper.urlopen(listingPage) 
    soup = BeautifulSoup(page, features="lxml")
    positionsPreTextScrape = []
    articleBodies = soup.find_all('article')
    tracking_number = counter
    for article in articleBodies:
        articleDivs = article.find_all('div')
        #logoDiv = articleDivs[0]
        bodyDiv = articleDivs[2]
        url = bodyDiv.find('a', href=True)
        bodyChildDivs = bodyDiv.find_all('div')
        company = bodyChildDivs[1]
        location = bodyDiv.find('li', class_='job-element__body__location')
        detailTable = bodyDiv.find('dl')
        detailsTitles = detailTable.find_all('dt')
        detailsValues = detailTable.find_all('dd')
        if url['href']:
            detailsText = []
            for detail in zip(detailsTitles, detailsValues):
                detailsText.append((detail[0].text, detail[1].text))
            urlString = "http://www.pnet.co.za"+url["href"]
            urlBase, parameters = urlString.split('?')
            newListing = JobListing(f"JL{str(tracking_number)}", urlBase, company.text.strip(), url.text.strip(), detailsText, location.text.strip())
            positionsPreTextScrape.append(newListing)
            tracking_number = tracking_number + 1
            logging.info(newListing)
    return positionsPreTextScrape, tracking_number

def testDetailedScrape(positionsPreTextScrape, session, delayValue):
    """ takes in the initial job listings that were scraped and generates the full job listings and docx files for each job listing """
    for position in positionsPreTextScrape:
        print(f"{position.id}:{position.url}")
        #scrape the position
        jobPage = position.url
        page = scraper.urlopen(jobPage) 
        jobSoup = BeautifulSoup(page, features="lxml")
        liquidDesignContainers = jobSoup.find_all('div', class_='listing-content__liquiddesign_container')
        description = jobSoup.find('div', class_='js-app-ld-ContentBlock')   
        mainContent = jobSoup.find('div', class_='listing__main-content')
        if description:
            elements = description.find_all(['h4','h3', 'h2', 'h1', 'p', 'li'])
            print(f"***description with {len(elements)} elements")
        elif mainContent:
            elements = mainContent.find_all(['h4','h3', 'h2', 'h1', 'p', 'li'])
            print(f"***mainContent with {len(elements)} elements")
        elif len(liquidDesignContainers) > 0:
            elements = []
            for liquidContainer in liquidDesignContainers:
                elements.extend(liquidContainer.find_all(['h4','h3', 'h2', 'h1', 'p', 'li']))
            print(f"***liquidDesign with {len(elements)} elements")
        else:
            print("***fail")
        time.sleep(delayValue)

def scrapeDetailedListings(positionsPreTextScrape, session, delayValue):
    """ takes in the initial job listings that were scraped and generates the full job listings and docx files for each job listing """
    for position in positionsPreTextScrape:
        document = Document()
        document.add_heading(position.title, level=1)
        document.add_heading(position.company, level=2)
        document.add_heading(position.location, level=3)
        detailString = ""
        for detail in position.details:
            detailString += f'{detail[0]} : {detail[1]} \n'
        document.add_heading(detailString, level=4)
        document.add_paragraph(position.url)
        print(position)
        #scrape the position
        jobPage = position.url
        page = scraper.urlopen(jobPage) 
        jobSoup = BeautifulSoup(page, features="lxml")
        #description = jobSoup.find_all('div', class_='listing-content__liquiddesign_container')
        liquidDesignContainers = jobSoup.find_all('div', class_='listing-content__liquiddesign_container')
        description = jobSoup.find('div', class_='js-app-ld-ContentBlock')   
        mainContent = jobSoup.find('div', class_='listing__main-content')
        # there are two distinct layouts, if the first one doesn't work, use the second one
        elements = []
        if description:
            elements.extend(description.find_all(['h4','h3', 'h2', 'h1', 'p', 'li']))
            print(f"***description with {len(elements)} elements")
        elif mainContent:
            elements.extend(mainContent.find_all(['h4','h3', 'h2', 'h1', 'p', 'li']))
            print(f"***mainContent with {len(elements)} elements")
        elif len(liquidDesignContainers) > 0:
            for liquidContainer in liquidDesignContainers:
                elements.extend(liquidContainer.find_all(['h4','h3', 'h2', 'h1', 'p', 'li']))
            print(f"***liquidDesign with {len(elements)} elements")
        else:
            print("***fail")
        for element in elements:
            if element.name == 'h4':
                document.add_heading(element.text.strip(), level=4)
                position.append_text_line(f'{element.text.strip()} \n')
            elif element.name == 'h3':
                document.add_heading(element.text.strip(), level=3)
                position.append_text_line(f'{element.text.strip()} \n')
            elif element.name == 'h2':
                document.add_heading(element.text.strip(), level=2)
                position.append_text_line(f'{element.text.strip()} \n')
            elif element.name == 'h1':
                document.add_heading(element.text.strip(), level=1)
                position.append_text_line(f'{element.text.strip()} \n')
            elif element.name == 'p':
                document.add_paragraph(element.text.strip())
                position.append_text_line(f'{element.text.strip()} \n')
            elif element.name == 'li':
                document.add_paragraph(element.text.strip(), style='List Bullet')
                position.append_text_line(f'* {element.text.strip()} \n')
            elif element.name == 'div':
                pass
                #document.add_paragraph(element.text.strip())
            elif element.name == 'strong':
                pass
                #document.add_paragraph(element.text.strip())
            else:
                document.add_paragraph(element.text.strip())
                position.append_text_line(f'{element.text.strip()} \n')
        logging.info(f"{position.id} scraped with detail")
        document.save(f"./Documents/{session}/{position.id}.docx")
        logging.info(f"{position.id} docx file created")
        time.sleep(delayValue)



# setting the session and session directory
session = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")
print(session)
os.makedirs(f"./Documents/{session}", exist_ok=True)
#os.mkdir(f"./Documents/{session}")
logging.basicConfig(level = logging.INFO, filename = f"./Documents/{session}/{time.strftime('my-%Y-%m-%d.log')}")

# pull the prelim listings from the advanced search pages
#advSearchUrls = ['https://www.pnet.co.za/5/job-search-detailed.html?searchTypeFrom=detailedSearch&searchOrigin=Detailed-Search_detailed-search&newsearch=1&Function=1025000&Function=1002000&Function=1004000&Function=1003002&Function=1003001&Function=1005002&Function=1005003&Function=1005001&Function=1013005&Function=1013004&Function=1003004&Function=1001000&Function=1007000&Function=1014000&Function=1018000&Function=1012000&Function=1008000&Function=1009000&Function=1011000&Function=1015000&Function=1006002&Function=1006001&Function=1003005&Function=1019000&Function=1013002&Function=1003003&Function=1003006&Function=1003007&Function=1016000&Function=1017000&Function=1006003&Function=1010000&Function=1013001&keyword=',
#'https://www.pnet.co.za/5/job-search-detailed.html?fu=1025000%2C1002000%2C1004000%2C1003002%2C1003001%2C1005002%2C1005003%2C1005001%2C1013005%2C1013004%2C1003004%2C1001000%2C1007000%2C1014000%2C1018000%2C1012000%2C1008000%2C1009000%2C1011000%2C1015000%2C1006002%2C1006001%2C1003005%2C1019000%2C1013002%2C1003003%2C1003006%2C1003007%2C1016000%2C1017000%2C1006003%2C1010000%2C1013001&of=25&suid=e8cd723a%2D6bc3%2D476d%2Da361%2D3b32743e50ae&an=paging%5Fnext&action=paging%5Fnext']

advSearchUrls = ['https://www.pnet.co.za/5/job-search-detailed.html?searchTypeFrom=detailedSearch&searchOrigin=Detailed-Search_detailed-search&newsearch=1&Function=1025000&keyword=']


PrelimScrapeListings = []
#scrapedListings, counter = scrapePrelimListingsNotClasses(1, advSearchUrls[0])
#PrelimScrapeListings.extend(scrapedListings)
#scrapedListings, counter = scrapePrelimListingsNotClasses(counter, advSearchUrls[1])
#PrelimScrapeListings.extend(scrapedListings)
counter = 1
for searchUrl in advSearchUrls:
    scrapedListings, counter = scrapePrelimListingsNotClasses(counter, searchUrl)
    PrelimScrapeListings.extend(scrapedListings)

# create the session directory and then pickle the prelim listings to the session
prelimScrapeFile = open(f'./Documents/{session}/prelimScrape', 'wb') # write binary
pickle.dump(PrelimScrapeListings, prelimScrapeFile)
prelimScrapeFile.close()

# do the secondary scrape where the text is added to the listings and the docx files are created
scrapeDetailedListings(PrelimScrapeListings, session, 2)
#testDetailedScrape(PrelimScrapeListings, session, 1)
secondaryScrapeFile = open(f'./Documents/{session}/secondaryScrape', 'wb') # write binary
pickle.dump(PrelimScrapeListings, secondaryScrapeFile) # the detailed scrape will have updated the prelimscrape listings
secondaryScrapeFile.close()

prelimReadfile = open(f'./Documents/{session}/prelimScrape', 'rb')
testPrelims = pickle.load(prelimReadfile)
prelimReadfile.close()
secondaryReadfile = open(f'./Documents/{session}/secondaryScrape', 'rb')
testSecondary = pickle.load(secondaryReadfile)
secondaryReadfile.close()

print(testPrelims[1])
print(testSecondary[1])

# From here you need to add in some of the parameters that you used to do the search and the scrape as a pickle
# then you need to do the app that reads the pickles and lets you code against the text